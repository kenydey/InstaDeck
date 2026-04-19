"""
文档解析模块 DocumentParser
支持 .txt, .md, .docx, .pdf 文件格式
将各种格式文档转换为纯文本大纲，保留表格结构
"""

import os
import re
from typing import Optional, List, Tuple
import warnings

# 尝试导入相关库，如果缺失则给出提示
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    warnings.warn("python-docx 未安装，将无法解析 .docx 文件。请运行：pip install python-docx")

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    warnings.warn("pdfplumber 未安装，将无法解析 .pdf 文件。请运行：pip install pdfplumber")

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    warnings.warn("PyMuPDF 未安装，将无法使用备用 PDF 解析器。请运行：pip install PyMuPDF")


class DocumentParser:
    """文档解析器，支持多种格式"""
    
    def __init__(self, prefer_pdfplumber=True):
        """
        初始化解析器
        
        Args:
            prefer_pdfplumber: 如果为 True，优先使用 pdfplumber 解析 PDF（表格提取更好）
                              如果为 False，使用 PyMuPDF（速度更快）
        """
        self.prefer_pdfplumber = prefer_pdfplumber
        
    def parse(self, file_path: str) -> str:
        """
        解析文档文件，返回纯文本大纲
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            解析后的纯文本内容
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在：{file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt':
            return self._parse_txt(file_path)
        elif ext == '.md':
            return self._parse_md(file_path)
        elif ext == '.docx':
            return self._parse_docx(file_path)
        elif ext == '.pdf':
            return self._parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式：{ext}。支持格式：.txt, .md, .docx, .pdf")
    
    def _parse_txt(self, file_path: str) -> str:
        """解析纯文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content
    
    def _parse_md(self, file_path: str) -> str:
        """解析 Markdown 文件，去除特殊符号但保留结构"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 保留标题、列表、表格等结构，去除纯格式符号
        # 将标题转换为纯文本（保留 # 符号作为层级指示）
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # 处理标题：保留 # 但去除多余空格
            if line.strip().startswith('#'):
                # 计算标题级别
                match = re.match(r'^(#+)\s*(.*)', line)
                if match:
                    level = len(match.group(1))
                    title = match.group(2)
                    cleaned_lines.append(f"{'#' * level} {title}")
                else:
                    cleaned_lines.append(line)
            # 处理表格：保留表格结构
            elif '|' in line and ('---' in line or '===' in line):
                # 表格分隔行，保留
                cleaned_lines.append(line)
            elif '|' in line:
                # 表格数据行，保留
                cleaned_lines.append(line)
            # 处理列表
            elif line.strip().startswith(('- ', '* ', '+ ', '1.', '2.', '3.')):
                cleaned_lines.append(line)
            else:
                # 普通文本，去除多余格式符号但保留内容
                cleaned_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)  # 去除链接
                cleaned_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', cleaned_line)  # 去除粗体
                cleaned_line = re.sub(r'\*([^*]+)\*', r'\1', cleaned_line)  # 去除斜体
                cleaned_line = re.sub(r'~~([^~]+)~~', r'\1', cleaned_line)  # 去除删除线
                cleaned_line = re.sub(r'`([^`]+)`', r'\1', cleaned_line)  # 去除内联代码
                cleaned_lines.append(cleaned_line)
        
        return '\n'.join(cleaned_lines)
    
    def _parse_docx(self, file_path: str) -> str:
        """解析 Word 文档"""
        if not HAS_DOCX:
            raise ImportError("请安装 python-docx：pip install python-docx")
        
        doc = docx.Document(file_path)
        content_parts = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)
        
        # 提取表格
        for table in doc.tables:
            table_text = self._docx_table_to_markdown(table)
            if table_text:
                content_parts.append(table_text)
        
        return '\n\n'.join(content_parts)
    
    def _docx_table_to_markdown(self, table) -> str:
        """将 Word 表格转换为 Markdown 格式"""
        if len(table.rows) == 0:
            return ""
        
        # 提取表格数据
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_data.append(cell_text)
            table_data.append(row_data)
        
        if not table_data:
            return ""
        
        # 转换为 Markdown 表格
        return self._list_to_markdown_table(table_data)
    
    def _parse_pdf(self, file_path: str) -> str:
        """解析 PDF 文档"""
        # 优先使用 pdfplumber（表格提取更好）
        if self.prefer_pdfplumber and HAS_PDFPLUMBER:
            return self._parse_pdf_with_pdfplumber(file_path)
        elif HAS_PYMUPDF:
            return self._parse_pdf_with_pymupdf(file_path)
        else:
            raise ImportError("请安装 pdfplumber 或 PyMuPDF：pip install pdfplumber 或 pip install PyMuPDF")
    
    def _parse_pdf_with_pdfplumber(self, file_path: str) -> str:
        """使用 pdfplumber 解析 PDF"""
        content_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # 提取文本
                text = page.extract_text()
                if text and text.strip():
                    content_parts.append(f"=== 第 {page_num} 页 ===")
                    content_parts.append(text.strip())
                
                # 提取表格
                tables = page.extract_tables()
                for table_num, table in enumerate(tables, 1):
                    if table:
                        table_text = self._list_to_markdown_table(table)
                        if table_text:
                            content_parts.append(f"--- 第 {page_num} 页表格 {table_num} ---")
                            content_parts.append(table_text)
        
        return '\n\n'.join(content_parts)
    
    def _parse_pdf_with_pymupdf(self, file_path: str) -> str:
        """使用 PyMuPDF 解析 PDF（备用方案）"""
        doc = fitz.open(file_path)
        content_parts = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            
            if text and text.strip():
                content_parts.append(f"=== 第 {page_num + 1} 页 ===")
                content_parts.append(text.strip())
        
        doc.close()
        return '\n\n'.join(content_parts)
    
    def _list_to_markdown_table(self, data: List[List[str]]) -> str:
        """
        将二维列表转换为 Markdown 表格
        
        Args:
            data: 二维列表，第一行通常是表头
            
        Returns:
            Markdown 格式的表格字符串
        """
        if not data:
            return ""
        
        # 计算每列的最大宽度
        col_widths = []
        for col_idx in range(len(data[0])):
            max_width = max(len(str(row[col_idx])) for row in data if col_idx < len(row))
            col_widths.append(max_width)
        
        # 构建表格
        lines = []
        
        # 表头
        header_cells = []
        for col_idx, cell in enumerate(data[0]):
            cell_str = str(cell)
            header_cells.append(cell_str.ljust(col_widths[col_idx]))
        lines.append('| ' + ' | '.join(header_cells) + ' |')
        
        # 分隔线
        separator_cells = []
        for width in col_widths:
            separator_cells.append('-' * width)
        lines.append('| ' + ' | '.join(separator_cells) + ' |')
        
        # 数据行（从第二行开始）
        for row in data[1:]:
            row_cells = []
            for col_idx, cell in enumerate(row):
                cell_str = str(cell)
                row_cells.append(cell_str.ljust(col_widths[col_idx]))
            lines.append('| ' + ' | '.join(row_cells) + ' |')
        
        return '\n'.join(lines)
    
    def parse_to_sections(self, file_path: str, min_section_length: int = 100) -> List[Tuple[str, str]]:
        """
        解析文档并分割为逻辑章节
        
        Args:
            file_path: 文档文件路径
            min_section_length: 最小章节长度（字符数）
            
        Returns:
            章节列表，每个元素为 (章节标题, 章节内容)
        """
        content = self.parse(file_path)
        
        # 根据标题分割章节
        sections = []
        current_section_title = "文档开头"
        current_section_content = []
        
        lines = content.split('\n')
        
        for line in lines:
            # 检测标题（支持 # 标题和类似 "1. Introduction" 的标题）
            is_heading = False
            heading_level = 0
            heading_text = ""
            
            # Markdown 风格标题
            md_match = re.match(r'^(#+)\s*(.+)$', line.strip())
            if md_match:
                is_heading = True
                heading_level = len(md_match.group(1))
                heading_text = md_match.group(2).strip()
            
            # 数字标题风格
            num_match = re.match(r'^(\d+(\.\d+)*)\s+(.+)$', line.strip())
            if not is_heading and num_match:
                is_heading = True
                heading_level = num_match.group(1).count('.') + 1
                heading_text = num_match.group(3).strip()
            
            # 大写字母标题风格
            caps_match = re.match(r'^([A-Z][A-Z\s]+)$', line.strip())
            if not is_heading and caps_match and len(line.strip()) < 50:
                is_heading = True
                heading_level = 2
                heading_text = line.strip()
            
            if is_heading:
                # 保存当前章节
                if current_section_content and len(' '.join(current_section_content)) > min_section_length:
                    sections.append((
                        current_section_title,
                        '\n'.join(current_section_content)
                    ))
                
                # 开始新章节
                current_section_title = heading_text
                current_section_content = [line]
            else:
                current_section_content.append(line)
        
        # 添加最后一个章节
        if current_section_content and len(' '.join(current_section_content)) > min_section_length:
            sections.append((
                current_section_title,
                '\n'.join(current_section_content)
            ))
        
        # 如果没有检测到章节，则将整个文档作为一个章节
        if not sections:
            sections.append(("全文", content))
        
        return sections


# 使用示例
if __name__ == '__main__':
    parser = DocumentParser()
    
    # 测试文件路径（请替换为实际文件路径）
    test_files = [
        "example.txt",
        "example.md",
        "example.docx",
        "example.pdf"
    ]
    
    for file_path in test_files:
        if os.path.exists(file_path):
            print(f"\n{'='*60}")
            print(f"解析文件：{file_path}")
            print('='*60)
            
            try:
                content = parser.parse(file_path)
                print(f"前1000个字符预览：")
                print(content[:1000] + "..." if len(content) > 1000 else content)
                
                # 测试章节分割
                sections = parser.parse_to_sections(file_path)
                print(f"\n检测到 {len(sections)} 个章节：")
                for i, (title, content) in enumerate(sections[:3]):  # 只显示前3个
                    print(f"  {i+1}. {title} ({len(content)} 字符)")
                
            except Exception as e:
                print(f"解析失败：{e}")
        else:
            print(f"文件不存在：{file_path}")